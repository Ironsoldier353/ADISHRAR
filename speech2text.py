from pydub import AudioSegment
import speech_recognition as sr
from docx import Document
import math

def split_audio(file_path, chunk_length_ms=60000):
    audio = AudioSegment.from_file(file_path)
    chunks = []
    for i in range(0, len(audio), chunk_length_ms):
        chunks.append(audio[i:i + chunk_length_ms])
    return chunks

# Example usage
#output_audio=r"C:\Users\SUSOVAN\iCloudDrive\Amrita\RBU\4th sem\Audio Recording\Geeti_Alekhyo\mp4_to_wav\Geeti_Alekhyo_Combined_All.wav"
output_audio=r'C:\Users\SUSOVAN\iCloudDrive\Amrita\RBU\4th sem\Audio Recording\Probondho\Probondho_Export.wav'
chunks = split_audio(output_audio)

# Save chunks as separate files
for i, chunk in enumerate(chunks):
    chunk.export(f"chunk_{i}.wav", format="wav")


def transcribe_google(audio_file):
    recognizer = sr.Recognizer()
    try:
        with sr.AudioFile(audio_file) as source:
            audio_data = recognizer.record(source)  # Read the audio file
            transcription = recognizer.recognize_google(audio_data, language='bn-IN')  # Change to 'en-US' for English
            return transcription
    except sr.UnknownValueError:
        return "Google Speech Recognition could not understand audio"
    except sr.RequestError as e:
        return f"Could not request results from Google Speech Recognition service; {e}"
    except ValueError as e:
        return f"Audio file error: {e}"

# Example: Process all chunks
doc = Document()
doc.add_heading("Audio Transcription", level=1)  # Add a heading

chunks = split_audio(output_audio)
for i, chunk in enumerate(chunks):
    chunk_path = f"chunk_{i}.wav"
    chunk.export(chunk_path, format="wav")
    transcription = transcribe_google(chunk_path)
    print(f"Transcription for chunk {i}:", transcription)
    doc.add_paragraph(transcription)  # Add the transcription as a paragraph
    



# Save the transcription to a .docx file
output_file_path = r'C:\Users\SUSOVAN\iCloudDrive\Amrita\RBU\4th sem\Audio Recording\Probondho\transcription.docx'
doc.save(output_file_path)
